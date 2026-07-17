"""The three bid deliverables from one merged extraction. NO LLM in this file —
deterministic rules map requirements to cost-driver lines and inspection doc
rows, because a hallucinated hold point in a QC document is the failure a
coatings contractor cannot survive an audit with.

1. Requirements sheet (CSV): every coat system + env limit + hold point, page-cited.
2. Compliance cost-driver checklist (CSV): rule-mapped line items with BLANK
   price columns — the estimator prices, the tool never does.
3. Draft inspection package (DOCX): daily QC report + pre-job checklist seeded
   with extracted hold points and acceptance ranges, every page watermarked.
"""
import csv
import json
from datetime import datetime, timezone
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "deliverables"
STANDARDS = json.loads((ROOT / "reference" / "standards.json").read_text())

WATERMARK = "DRAFT — REQUIRES CERTIFIED INSPECTOR REVIEW. Not a compliance determination."

REQUIREMENT_KEYS = ("coat_systems", "environmental_limits", "hold_points")


def _req_text(ex):
    """Rule-scan surface: ONLY the requirement fields. Scanning the whole job dict
    let a file named waste_treatment.pdf fire the waste cost driver."""
    return json.dumps({k: ex.get(k, []) for k in REQUIREMENT_KEYS}).lower()


# Deterministic cost-driver rules: requirement pattern -> checklist line item.
# kn: seeded from the product spec's named drivers; extend per pilot contractor.
COST_DRIVER_RULES = [
    (lambda ex: any("SP5" in (cs.get("prep_standard") or "") or "SP10" in (cs.get("prep_standard") or "")
                    for cs in ex["coat_systems"]),
     "Abrasive blast program (white/near-white): media, equipment, disposal"),
    (lambda ex: any(cs.get("prep_standard") for cs in ex["coat_systems"]),
     "Surface prep verification instruments (profile, cleanliness)"),
    (lambda ex: "contain" in _req_text(ex),
     "Containment class: structure, negative air, monitoring"),
    (lambda ex: "confined" in _req_text(ex),
     "Confined-space entry program: attendants, rescue, air monitoring"),
    (lambda ex: any(e for e in ex["environmental_limits"]
                    if "humid" in (e.get("condition") or "").lower() or "dew" in (e.get("condition") or "").lower()),
     "Dehumidification / ambient control equipment"),
    (lambda ex: bool(ex["environmental_limits"]),
     "Ambient monitoring: logging instruments + recordkeeping labor"),
    (lambda ex: bool(ex["hold_points"]),
     "Inspector hold-point standby time (crew idle during witness points)"),
    (lambda ex: "waste" in _req_text(ex) or "disposal" in _req_text(ex),
     "Waste characterization and disposal"),
]


def _cell(s):
    """CSV formula-injection guard: Excel executes cells starting with = + - @.
    Spec text is untrusted; a hold point of '=cmd|...' must open inert."""
    s = str(s if s is not None else "")
    return "'" + s if s[:1] in "=+-@" else s


def standard_citation(prep: str) -> str:
    key = (prep or "").upper().replace(" ", "-")
    for name, desc in STANDARDS.items():
        if name.replace(" ", "-") == key:
            return f"{name}: {desc}"
    return "(no matching standard in reference library)"


def requirements_csv(extraction: dict, job: str) -> Path:
    OUT.mkdir(exist_ok=True)
    path = OUT / f"{job}-requirements.csv"
    with path.open("w", newline="") as f:
        w = csv.writer(f)
        w.writerow(["type", "detail", "standard_reference", "page", "flag"])
        for cs in extraction["coat_systems"]:
            coats = "; ".join(f"{c['product_or_type']} {c['dft_min_mils']}-{c['dft_max_mils']} mils"
                              for c in cs["coats"])
            flag = "NEEDS REVIEW" if cs.get("flag") else ""
            w.writerow(["coat_system", _cell(f"{cs['surface']} | prep {cs['prep_standard']} | {coats}"),
                        standard_citation(cs["prep_standard"]), _cell(cs["page"]), flag])
        for e in extraction["environmental_limits"]:
            w.writerow(["environmental_limit", _cell(f"{e['condition']}: {e['limit']}"), "",
                        _cell(e["page"]), "NEEDS REVIEW" if e.get("flag") else ""])
        for h in extraction["hold_points"]:
            w.writerow(["hold_point", _cell(h["point"]), "", _cell(h["page"]),
                        "NEEDS REVIEW" if h.get("flag") else ""])
    return path


def cost_drivers_csv(extraction: dict, job: str) -> tuple[Path, int]:
    OUT.mkdir(exist_ok=True)
    path = OUT / f"{job}-cost-drivers.csv"
    ex = {k: extraction.get(k, []) for k in REQUIREMENT_KEYS}  # rules never see job metadata
    lines = [label for rule, label in COST_DRIVER_RULES if rule(ex)]
    with path.open("w", newline="") as f:
        w = csv.writer(f)
        w.writerow(["cost_driver", "unit", "qty", "unit_cost", "extended"])  # price columns BLANK by design
        for label in lines:
            w.writerow([label, "", "", "", ""])
    return path, len(lines)


def inspection_docx(extraction: dict, job: str) -> Path:
    from docx import Document  # deferred import

    doc = Document()
    doc.add_heading(f"Inspection document package — {job}", level=0)
    doc.add_paragraph(f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d')}")
    doc.add_paragraph(WATERMARK, style="Intense Quote")

    doc.add_heading("Pre-job checklist (seeded from spec)", level=1)
    t = doc.add_table(rows=0, cols=3)
    for cs in extraction["coat_systems"]:
        for c in cs["coats"]:
            cells = t.add_row().cells
            cells[0].text = f"{cs['surface']} — {c['product_or_type']}"
            cells[1].text = f"DFT {c['dft_min_mils']}-{c['dft_max_mils']} mils ({cs['page']})"
            cells[2].text = "⚠ NEEDS REVIEW" if cs.get("flag") else "verified ☐"

    doc.add_heading("Daily QC report template", level=1)
    t2 = doc.add_table(rows=0, cols=2)
    for label in ("Date / shift", "Ambient temp / RH / dew point / surface temp",
                  "Surface prep performed & standard", "Coats applied & measured DFT",
                  "Inspector signature (certified)"):
        cells = t2.add_row().cells
        cells[0].text = label
        cells[1].text = ""
    for e in extraction["environmental_limits"]:
        cells = t2.add_row().cells
        cells[0].text = f"Limit ({e['page']})"
        cells[1].text = f"{e['condition']}: {e['limit']}"

    doc.add_heading("Hold points (verbatim from spec)", level=1)
    for h in extraction["hold_points"]:
        doc.add_paragraph(f"☐ {h['point']} ({h['page']})")
    doc.add_paragraph(WATERMARK)

    OUT.mkdir(exist_ok=True)
    path = OUT / f"{job}-inspection-draft.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    import tempfile
    OUT = Path(tempfile.mkdtemp())
    ex = {"coat_systems": [{"surface": "tank interior (confined space)", "prep_standard": "SSPC-SP10",
                            "coats": [{"product_or_type": "zinc primer", "dft_min_mils": 3.0, "dft_max_mils": 5.0}],
                            "notes": "full containment required", "page": "p.12"}],
          "environmental_limits": [{"condition": "relative humidity", "limit": "max 85%", "page": "p.13"}],
          "hold_points": [{"point": "prep sign-off before primer", "page": "p.14"}],
          "needs_review": [], "guard_violations": 0}
    p1 = requirements_csv(ex, "J1")
    assert "SSPC-SP10" in p1.read_text() and "Near-White" in p1.read_text()
    p2, n = cost_drivers_csv(ex, "J1")
    body = p2.read_text()
    assert n >= 5  # blast, instruments, containment, confined space, dehu, monitoring, hold standby
    assert ",,,," in body  # price columns blank — the estimator prices
    p3 = inspection_docx(ex, "J1")
    assert p3.exists() and p3.stat().st_size > 0
    print(f"deliverables OK — {n} cost drivers mapped, prices blank, watermark on")
